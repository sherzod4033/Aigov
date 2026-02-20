from docx import Document
from docx.shared import Pt

document = Document()

# Title
document.add_heading('Справочник Налогоплательщика (Тестовый)', 0)

# Intro
p = document.add_paragraph('Этот документ создан для тестирования системы AndozAI. ')
p.add_run('Он содержит базовые примеры статей и правил.').bold = True

# Article 1
document.add_heading('Статья 1. Общие положения', level=1)
document.add_paragraph(
    'Налог — это обязательный, индивидуально безвозмездный платеж, взимаемый с организаций и физических лиц '
    'в форме отчуждения принадлежащих им на праве собственности денежных средств.'
)

# Article 2
document.add_heading('Статья 2. Налог на прибыль', level=1)
document.add_paragraph(
    '1. Объектом налогообложения признается прибыль, полученная налогоплательщиком.\n'
    '2. Стандартная ставка налога на прибыль устанавливается в размере 15%.\n'
    '3. Для IT-компаний действует льготная ставка в размере 7.5% при условии аккредитации.'
)

# Article 3
document.add_heading('Статья 3. Сроки уплаты', level=1)
document.add_paragraph(
    'Налоговый период составляет один календарный год. Авансовые платежи должны быть внесены '
    'не позднее 28 числа месяца, следующего за отчетным периодом.'
)

# FAQ Section for testing Retrieval
document.add_heading('Часто задаваемые вопросы (FAQ)', level=1)
document.add_paragraph('Вопрос: Какие льготы есть для стартапов?', style='List Bullet')
document.add_paragraph('Ответ: Стартапы освобождаются от проверок на 2 года.', style='List Bullet')

document.add_paragraph('Вопрос: Какой штраф за просрочку декларации?', style='List Bullet')
document.add_paragraph('Ответ: Штраф составляет 5% от суммы налога за каждый месяц просрочки.', style='List Bullet')

# Save
file_path = 'test_tax_code.docx'
document.save(file_path)
print(f"Document saved to {file_path}")
